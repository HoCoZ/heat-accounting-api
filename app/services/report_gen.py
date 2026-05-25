from datetime import datetime
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ReportGenerator:

    @staticmethod
    def generate_balance_act(
        period_start: datetime,
        period_end: datetime,
        total_supplied: float,
        total_consumed: float,
        loss: float,
        loss_percent: float,
        is_balanced: bool,
        consumer_details: list,
    ) -> bytes:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(1.5)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("АКТ ПРИЕМА-ПЕРЕДАЧИ ТЕПЛОВОЙ ЭНЕРГИИ")
        run.bold = True
        run.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"за период с {period_start.strftime('%d.%m.%Y')} по {period_end.strftime('%d.%m.%Y')}")
        run.font.size = Pt(12)

        doc.add_paragraph()

        doc.add_paragraph(f"Общий отпуск тепла: {total_supplied:.2f} Гкал")
        doc.add_paragraph(f"Общее потребление: {total_consumed:.2f} Гкал")
        doc.add_paragraph(f"Потери в сетях: {loss:.2f} Гкал ({loss_percent:.1f}%)")

        status = "СОШЕЛСЯ" if is_balanced else "НЕ СОШЕЛСЯ"
        p = doc.add_paragraph()
        run = p.add_run(f"Баланс: {status}")
        run.bold = True

        if not is_balanced:
            doc.add_paragraph("Причина: расхождение превышает допустимый порог в 5%.")

        doc.add_paragraph()

        if consumer_details:
            doc.add_paragraph("Детализация по потребителям:", style="List Bullet")
            for c in consumer_details:
                doc.add_paragraph(f"{c['name']}: {c['consumed_gcal']:.2f} Гкал", style="List Bullet 2")

        file_path = f"act_{period_start.strftime('%Y%m%d')}_{period_end.strftime('%Y%m%d')}.docx"
        doc.save(file_path)
        with open(file_path, "rb") as f:
            data = f.read()
        os.remove(file_path)
        return data
